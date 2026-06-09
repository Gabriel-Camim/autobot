from __future__ import annotations

import json
import re
import shutil
from pathlib import Path

import yaml
from docx import Document


ROOT = Path(__file__).resolve().parents[2]
UPDATED = "2026-06-09"
KNOWLEDGE = ROOT / "knowledge"
MATERIALS = ROOT / "materials" / "recruiter-pack"
DGE_ROOT = Path(r"C:\Users\diagn\OneDrive\Documents\dge-metacognition-db\dge-2.0")
CV_DOCX = Path(r"C:\Users\diagn\Downloads\Curriculo_Gabriel_Camim_Dados_IA_Automacao (2).docx")
EBOOK_JSON = Path(r"C:\Users\diagn\Downloads\Base ebook generator (19).json")


def fm(title: str, category: str, tags: list[str], priority: int, summary: str, **extra) -> str:
    meta = {
        "title": title,
        "category": category,
        "tags": tags,
        "visibility": "public",
        "priority": priority,
        "updated_at": UPDATED,
        "summary": summary,
    }
    meta.update(extra)
    return "---\n" + yaml.safe_dump(meta, allow_unicode=True, sort_keys=False).strip() + "\n---\n\n"


def write_md(path: str | Path, title: str, category: str, tags: list[str], priority: int, summary: str, body: str, **extra) -> None:
    target = ROOT / path if isinstance(path, str) else path
    target.parent.mkdir(parents=True, exist_ok=True)
    target.write_text(fm(title, category, tags, priority, summary, **extra) + body.strip() + "\n", encoding="utf-8", newline="\n")


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " | ") for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def bulletize(items: list[str]) -> str:
    return "\n".join(f"- {item}" for item in items)


cv_text = docx_text(CV_DOCX)
ebook_data = json.loads(EBOOK_JSON.read_text(encoding="utf-8"))
ebook_nodes = [node.get("name", "") for node in ebook_data.get("nodes", []) if node.get("name")]

docs: list[tuple[str, str, str, list[str], int, str, str]] = [
    (
        "knowledge/gabriel/perfil.md",
        "Perfil profissional de Gabriel Camim Santos",
        "gabriel",
        ["perfil", "ia", "dados", "automacao", "revops", "arquitetura"],
        1,
        "Perfil profissional consolidado a partir do currículo de Gabriel.",
        """
# Perfil profissional

Sou Gabriel Camim Santos, profissional de tecnologia em transição e consolidação para Data Science, IA generativa, automações, RevOps e arquitetura de sistemas. Tenho formação técnica em Informática, experiência prática em suporte a sistemas críticos e um histórico forte de projetos autorais em IA, dados, e-commerce, automação editorial e medicina veterinária computacional.

Minha trajetória combina dois lados que eu gosto de manter juntos: vivência operacional real e construção de sistemas. Já trabalhei em ambientes onde sistema parado vira impacto imediato para a operação, como frente de caixa, ERP, infraestrutura hospitalar e suporte N2. Ao mesmo tempo, venho desenvolvendo projetos autorais que exigem modelagem de dados, arquitetura modular, APIs, workflows, LLMs, prompts, RAG, bancos relacionais, automações e visão de produto.

Em entrevista, eu me posiciono como alguém que aprende rápido, estrutura problema antes de sair codando e tenta transformar ambiguidade em fluxo, dados e decisão. Gosto de construir sistemas que não sejam apenas bonitos ou demonstrativos, mas que organizem conhecimento, reduzam atrito operacional e criem inteligência aplicável ao negócio.

Meus interesses principais hoje são Data Science, IA aplicada, automação, sistemas internos, inteligência operacional, RevOps, produtos baseados em dados e interfaces que tornem raciocínios complexos mais claros para pessoas reais.
""",
    ),
    (
        "knowledge/gabriel/curiosidades.md",
        "Curiosidades pessoais",
        "gabriel",
        ["curiosidades", "personalidade", "historia-pessoal"],
        2,
        "Curiosidades pessoais que ajudam a mostrar o modo de aprender e a identidade do Gabriel.",
        """
# Curiosidades pessoais

Algumas curiosidades ajudam a explicar meu jeito de aprender e de me relacionar com conhecimento:

- Aprendi a ler e escrever com 4 anos de idade, antes mesmo de ir ao jardim de infância.
- Passei na ETEC com a segunda maior nota da instituição.
- Consegui bolsa 100% pelo ProUni para estudar Medicina Veterinária sem abrir um livro de estudo para aquela prova.
- Tenho 14 tatuagens, sendo 12 de animais.

Esses pontos não entram como troféu vazio. Para mim, eles contam uma história de autonomia cognitiva, curiosidade, relação muito forte com animais e facilidade de aprender por conta própria quando encontro um objetivo que faz sentido.
""",
    ),
    (
        "knowledge/gabriel/objetivos-pessoais.md",
        "Objetivos pessoais",
        "gabriel",
        ["objetivos", "data-science", "produto", "impacto"],
        1,
        "Objetivos pessoais e profissionais declarados pelo Gabriel.",
        """
# Objetivos pessoais

Meus objetivos pessoais e profissionais atuais são claros:

- Me formar em Data Science.
- Criar alguma tecnologia nova, ou ao menos um sistema/app que seja globalmente utilizado.

O que me move é a possibilidade de construir algo que saia do campo da ideia e vire infraestrutura útil para muita gente. Eu gosto de sistemas que capturam complexidade, organizam informação e ajudam pessoas ou empresas a decidir melhor.
""",
    ),
    (
        "knowledge/trajetoria/formacao-academica.md",
        "Formação acadêmica e cursos",
        "trajetoria",
        ["formacao", "fiap", "unisa", "etec", "idiomas"],
        1,
        "Formação acadêmica, cursos e idiomas do Gabriel.",
        """
# Formação acadêmica e cursos

## Data Science - FIAP

Tenho previsão de cursar Data Science na FIAP entre julho de 2026 e dezembro de 2028. Essa formação está alinhada com meu objetivo de consolidar minha atuação em dados, IA aplicada, modelagem, analytics e sistemas inteligentes.

## Medicina Veterinária - UNISA

Cursei Medicina Veterinária na UNISA entre fevereiro de 2023 e setembro de 2025, em trajetória incompleta. Mesmo sem concluir o curso, esse período virou uma base aplicada importante para projetos de IA, dados clínicos, classificação de doenças, modelagem de conhecimento, terminologias controladas e sistemas de apoio à decisão.

## Ensino Médio Técnico em Informática - ETEC Adhemar Batista Heméritas

Concluí o Ensino Médio Técnico em Informática entre 2018 e 2020. A base técnica incluiu lógica de programação, análise de sistemas, desenvolvimento de software, programação orientada a objetos, banco de dados, redes, manutenção, sistemas operacionais, mobilidade, empreendedorismo e inovação.

## Cursos complementares

- UX Design de A a Z - Udemy.
- UI Design com Figma: do zero ao especialista - Udemy.

## Idiomas

Inglês com leitura fluente, escrita avançada e conversação avançada.
""",
    ),
    (
        "knowledge/trajetoria/linha-do-tempo.md",
        "Linha do tempo profissional",
        "trajetoria",
        ["linha-do-tempo", "carreira", "projetos"],
        1,
        "Linha do tempo resumida da trajetória de Gabriel.",
        """
# Linha do tempo

## 2018-2020: base técnica em Informática

Construí minha base inicial em tecnologia na ETEC Adhemar Batista Heméritas, com programação, banco de dados, redes, análise de sistemas, desenvolvimento e manutenção.

## 2021-2023: suporte em ambiente hospitalar crítico

Atuei pela Spread Tecnologia no Hospital Israelita Albert Einstein, primeiro em rollout de parque computacional e depois em formatação, configuração, instalação de sistemas, Field Service e Service Desk N2 remoto. Esse período me deu disciplina operacional, noção de criticidade e contato com infraestrutura real.

## 2023-2025: Medicina Veterinária e IA aplicada à saúde animal

Na UNISA, conectei Medicina Veterinária com ciência de dados, modelagem relacional, classificação multieixo e IA aplicada. Desse ciclo nasceram DiagnósIA, VetDex, VDC e Córtex Digital.

## 2025: ERP, PDV e operação de franquias

Atuei como Analista de Suporte Jr na Software Show / Cacau Show, prestando suporte aos sistemas Master Retail PDV e Master Retail Desktop, com foco em estabilidade de loja, emissão fiscal, pagamentos, retaguarda e integrações.

## 2026 em diante: projetos autorais em dados, IA e automação

A partir de 2026, passei a consolidar projetos autorais como DGE, Ebook Generator e Autobot, combinando IA generativa, arquitetura de sistemas, automações, dados e visão de produto.
""",
    ),
    (
        "knowledge/skills/hard-skills.md",
        "Hard skills",
        "skills",
        ["hard-skills", "python", "sql", "llm", "automacao", "dados"],
        1,
        "Competências técnicas consolidadas a partir do currículo.",
        """
# Hard skills

## IA generativa e LLMs

OpenAI, Claude, LangChain, engenharia de prompts, agentes, structured outputs, tool calling, RAG, guardrails, avaliação de respostas, memória, orquestração de modelos e desenvolvimento assistido por IA com Codex e Abacus AI.

## Dados e analytics

SQL, PostgreSQL, modelagem relacional, schemas, KPIs, projeções, forecast/reforecast, governança de dados, lineage, rastreabilidade, documentação de métricas, análise operacional, Power BI e Superset.

## Automação e integrações

n8n, APIs REST, webhooks, JSON, OpenAI API, Gamma API, parsing de respostas, tratamento de erros, pipelines de dados textuais, integrações entre ferramentas e workflows operacionais.

## Desenvolvimento e arquitetura

JavaScript, Python, HTML/CSS, GitHub, Visual Studio Code, Vercel, Docker, backend, APIs internas, arquitetura modular, autenticação, permissões, CI/CD e banco multitenant.

## Marketing digital e RevOps

CRM, ICP, funil comercial, B2B/B2C, lead scoring, automações de captação, dados para conversão de receita, análise de campanhas, product brain, checkout de pagamentos e inteligência de crescimento.
""",
    ),
    (
        "knowledge/skills/soft-skills.md",
        "Soft skills",
        "skills",
        ["soft-skills", "aprendizado", "comunicacao", "pensamento-analitico"],
        1,
        "Competências comportamentais que aparecem na trajetória e nos projetos.",
        """
# Soft skills

- Aprendizado autodidata e rápido, com forte capacidade de mergulhar em domínios novos.
- Pensamento analítico para decompor problemas grandes em entidades, fluxos, regras e decisões.
- Comunicação técnica com preocupação de traduzir complexidade para pessoas de negócio.
- Senso de dono em projetos autorais, sustentando documentação, arquitetura e evolução contínua.
- Atenção a rastreabilidade, governança e limites de automação em sistemas que apoiam decisão.
- Criatividade aplicada: transformar uma ideia abstrata em workflow, banco, interface, prompt, API ou documentação executável.
""",
    ),
    (
        "knowledge/experiencia/experiencias-profissionais.md",
        "Experiência profissional",
        "experiencia",
        ["experiencia", "suporte", "erp", "pdv", "hospital", "automacao"],
        1,
        "Experiências profissionais formais do Gabriel.",
        """
# Experiência profissional

## Analista de Suporte Jr - Software Show / Cacau Show | Jan/2025 - Jan/2025

Atuei no suporte aos sistemas Master Retail PDV e Master Retail Desktop utilizados por franquias da Cacau Show, com foco em estabilidade operacional, emissão fiscal, pagamentos e continuidade de loja.

Principais frentes:

- Suporte a ERP/retaguarda e frente de caixa.
- Apoio a rotinas administrativas, cadastros, controle de loja, vendas, pagamentos e operações fiscais.
- Diagnóstico de falhas em produção, reprodução de problemas, chamados, incidentes e integrações.
- Contato com TEF, SAT, pinpads, impressoras térmicas, totens e chamadores de senha.

## Analista de Suporte Jr - Spread Tecnologia / Hospital Israelita Albert Einstein | Ago/2021 - Fev/2023

Atuei em ambiente hospitalar de alta criticidade, começando com rollout de parque computacional e evoluindo para formatação, configuração, instalação de sistemas e preparação de estações Windows.

Principais frentes:

- Field Service e Service Desk N2 remoto.
- Suporte a usuários, hardware, software, sistemas corporativos e infraestrutura.
- Uso de Active Directory, SCCM, Cisco ACS, Cisco ISE e ITSM.
- Criação de automação em `.bat` para instalações e configurações padronizadas, posteriormente adotada pela equipe de suporte para acelerar o fluxo operacional.

Essas experiências formaram minha visão prática: tecnologia precisa funcionar em produção, com usuário real, pressão real e continuidade operacional.
""",
    ),
    (
        "knowledge/projetos/dge/overview.md",
        "DGE - Data Growth Engine",
        "projetos",
        ["dge", "data-growth-engine", "ecommerce", "inteligencia-operacional", "galpao"],
        1,
        "Visão geral da DGE como camada de inteligência operacional auditável.",
        """
# DGE - Data Growth Engine

A DGE é um sistema autoral para governança e implantação de e-commerce orientado por dados. A ideia central é estruturar uma camada de inteligência operacional auditável para conectar marketplaces, Bling, ERP, produtos, SKUs, pedidos, estoque, preços, compras, frete, canais próprios, fulfillment, BI e decisões de crescimento.

O projeto começou como tese de crescimento digital e evoluiu para um backbone operacional: regras, validações, contratos de dados, estados, cálculos, auditoria, permissões e módulos capazes de transformar fatos da operação em indicadores, projeções, gargalos, variâncias, impacto projetivo e reforecast governado.

A frase que resume a DGE: dados reais da operação entram por uma camada governada, viram KPIs, alimentam fórmulas, preservam snapshots, geram lineage, mostram variâncias, apontam gargalos e sustentam decisões humanas sem apagar histórico.

## O que a DGE não é

A DGE não é apenas um site, dashboard, ERP, BI, automação ou IA mágica. Ela usa elementos de todos esses mundos, mas seu diferencial é organizar a decisão: origem do dado, qualidade, fórmula afetada, versão do plano, impacto possível e fluxo de aprovação.

## Papel da IA

A IA futura na DGE deve atuar como camada de recomendação e explicação com contexto rastreável. Ela pode priorizar exceções, sugerir playbooks e explicar causas prováveis, mas não deve aprovar compra, estoque, cupom, cobrança, reforecast ou decisão fiscal sozinha.
""",
    ),
    (
        "knowledge/projetos/dge/inteligencia-operacional.md",
        "DGE 2.0 - Inteligência operacional auditável",
        "projetos",
        ["dge", "kpi", "reforecast", "lineage", "auditabilidade"],
        1,
        "Síntese executiva da evolução da DGE 2.0 para inteligência operacional auditável.",
        """
# DGE 2.0 - Inteligência operacional auditável

A DGE 2.0 mudou de natureza: deixou de ser só uma demonstração de potencial para começar a estruturar uma camada de inteligência operacional auditável. Em vez de responder apenas \"o que aconteceu?\", a DGE busca responder de onde veio o número, qual KPI ele altera, qual fórmula foi pressionada, qual projeção ficou em risco, qual gargalo apareceu, qual decisão precisa ser tomada e qual versão do plano deve ser preservada.

## Cadeia de decisão

A cadeia começa com um fato operacional: venda, pedido, estoque, frete, comissão, campanha, custo ou gargalo. Esse fato entra em Data Intake com origem, status, tipo, período e destino. Depois pode virar KPI, alimentar fórmula, gerar variância, produzir impacto projetivo e, se houver evidência e aprovação, sustentar reforecast oficial.

## Componentes-chave

- Data Intake: porta de entrada governada para fatos operacionais.
- KPI Intelligence: transforma entradas em indicadores auditáveis.
- Formula Registry: registra regras de cálculo, variáveis e versões.
- Snapshots: preservam o estado de dados, projeções e leituras.
- Lineage: mostra origem, cálculo, versão e revisão.
- Projection Impact Analyzer: calcula impacto possível sem alterar versão oficial.
- Official Reforecast: cria nova versão governada da projeção quando há evidência.

## Galpão como contexto

Nos materiais enviados, a DGE aparece aplicada ao contexto da Galpão, conectando e-commerce próprio, marketplaces, Bling, estoque, hubs, CDD, logística, BI, exceções e futuras integrações. Um exemplo apresentado envolve comportamento capturado de usuários, leitura de demanda reprimida, estoque disponível e relatório estruturado para o time decidir ação comercial. A inteligência apoia; a decisão continua humana.
""",
    ),
    (
        "knowledge/projetos/dge/arquitetura.md",
        "Arquitetura da DGE",
        "projetos",
        ["dge", "arquitetura", "postgresql", "multitenant", "api", "backend"],
        1,
        "Resumo técnico da arquitetura da DGE e dos módulos documentados.",
        """
# Arquitetura da DGE

A DGE separa ERP Core da camada de inteligência operacional. O ERP registra cadastros, estoque e operação; a DGE cruza esses dados com KPIs, projeções, fórmulas, reforecast, canais, logística, compras, exceções e governança.

## Pilares técnicos

- Backend modular com contratos de domínio.
- Banco relacional em PostgreSQL com modelagem de schemas.
- Estrutura multitenant, autenticação e permissões.
- APIs internas para módulos de produto, SKU, pedidos, canais, compras, estoque, frete, BI e exceções.
- Registros de lineage, snapshots, versões e estados.
- Smokes/checks para validar módulos e contratos.
- Roadmap Debt Register para diferenciar capacidade ativa, parcial, planejada e bloqueada.

## Módulos documentados na base DGE 2.0

A base de referência inclui documentação para Activation and KPI History, Adaptive Projection Engine, After Sales Claims Backbone, Automation Backbone, BI Semantic Layer, Bling ERP Connector, Channel Intelligence, Commerce Operational Events, Commerce Orders and Identity, Core Schema, Data Contracts, ERP Intelligence Integration, Event Registry, Forecast Reconciliation, Freight Economics, Fulfillment Control Tower, Inventory Capacity, KPI Intelligence, Marketplace Fulfillment Pools, Metric and Formula Registry, ML/Fine-tuning, Module Registry, Operational Cockpit, Operational Command Center, Operational Governance, Operational Nervous System, Projection Engine, Projection Impact Analyzer, Tenant Access and Responsibility e contratos de integração.

## Decisão arquitetural importante

A DGE evita esconder regra de negócio no frontend. A interface é uma superfície futura; a base real é o backend, onde vivem validações, contratos, cálculos, estados e auditoria.
""",
    ),
    (
        "knowledge/projetos/veterinaria/overview.md",
        "Projetos veterinários com IA e dados",
        "projetos",
        ["veterinaria", "ia", "dados-clinicos", "classificacao", "vetdex", "vdc"],
        1,
        "Visão geral dos projetos de IA, dados e classificação em Medicina Veterinária.",
        """
# Projetos veterinários com IA e dados

Durante a graduação em Medicina Veterinária, desenvolvi projetos que conectavam saúde animal, ciência de dados, classificação de doenças, modelagem relacional, IA aplicada e sistemas de apoio à decisão.

Os principais nomes desse ciclo são DiagnósIA, VetDex, VDC e Córtex Digital. Eles partem de uma dor comum: a Medicina Veterinária ainda convive com registros despadronizados, vocabulários heterogêneos, pouca interoperabilidade e dificuldade de organizar conhecimento clínico de forma computacional.

O objetivo desses projetos era transformar conhecimento veterinário em estrutura: entidades nosológicas, eixos clínicos, modificadores, metadados, vocabulários controlados, codificação alfanumérica, filtros multieixo, modelagem relacional e base para RAG, busca semântica, fine-tuning e modelos preditivos.
""",
    ),
    (
        "knowledge/projetos/veterinaria/vdc.md",
        "VDC - Veterinary Disease Classification",
        "projetos",
        ["vdc", "classificacao", "doencas", "multieixo", "medicina-veterinaria"],
        1,
        "Classificação multieixo de doenças veterinárias proposta por Gabriel.",
        """
# VDC - Veterinary Disease Classification

O VDC é uma proposta de classificação técnico-científica de doenças veterinárias baseada em modelo multieixo e codificação alfanumérica hierárquica. A Fase 1 foi desenhada para o Capítulo 1, dedicado a doenças infecciosas caninas.

## Problema

Na Medicina Veterinária, ainda há ausência de uma classificação técnico-científica padronizada e semanticamente estruturada comparável à maturidade de sistemas como CID-11 e SNOMED-CT na medicina humana. Essa lacuna dificulta interoperabilidade, prontuários informatizados, pesquisa clínica, vigilância epidemiológica e ensino estruturado.

## Proposta

O VDC organiza doenças por níveis, eixos clínicos e modificadores. Os eixos previstos incluem etiologia, fisiopatologia, manifestações clínicas, anatomopatologia, diagnóstico e conduta terapêutica. Modificadores representam variações relevantes, como predisposição racial e contextos específicos de manifestação patológica.

## Entrega piloto

A proposta previa uma aplicação piloto com o Capítulo 1 preenchido por dez doenças infecciosas caninas, validada em ambiente institucional, auditável, documentada e preparada para expansão futura.
""",
    ),
    (
        "knowledge/projetos/veterinaria/vetdex.md",
        "VetDex",
        "projetos",
        ["vetdex", "banco-relacional", "postgresql", "filtros", "classificacao"],
        1,
        "Sistema computacional de suporte ao VDC.",
        """
# VetDex

VetDex é o sistema computacional de suporte ao VDC. Enquanto o VDC é a classificação, o VetDex é a infraestrutura que operacionaliza a classificação por meio de banco de dados relacional, vocabulários controlados, filtros de consulta multieixo, versionamento supervisionado e rotinas de automação.

## Funcionalidades previstas

- Modelagem entidade-relacionamento para doenças, eixos, códigos, modificadores e vocabulários.
- Banco relacional compatível com SQLite/PostgreSQL.
- Filtros multieixo e consultas compostas.
- Scripts de automação e versionamento.
- Documentação técnica e manuais de uso.
- Validação funcional em ambiente acadêmico-clínico.

## Por que importa

O VetDex transforma uma classificação teórica em sistema consultável, auditável e expansível. Ele serve como base para prontuários, ensino, pesquisa, RAG, busca semântica e futuras aplicações de IA em saúde animal.
""",
    ),
    (
        "knowledge/projetos/veterinaria/diagnosia.md",
        "DiagnósIA",
        "projetos",
        ["diagnosia", "ia", "apoio-a-decisao", "veterinaria"],
        2,
        "Projeto de IA aplicada à organização e apoio ao raciocínio diagnóstico veterinário.",
        """
# DiagnósIA

DiagnósIA é um projeto de IA aplicada à Medicina Veterinária, voltado para apoio ao raciocínio diagnóstico, organização de conhecimento clínico e uso de dados estruturados para orientar investigação, diferenciais e hipóteses.

O projeto se conecta ao mesmo domínio do VetDex/VDC: transformar conhecimento veterinário em estrutura computável. A intenção não é substituir o médico-veterinário, mas criar apoio inteligente, rastreável e baseado em evidências para estudo, triagem, raciocínio clínico e organização de dados.
""",
    ),
    (
        "knowledge/projetos/veterinaria/cortex-digital.md",
        "Córtex Digital",
        "projetos",
        ["cortex-digital", "ia", "conhecimento", "veterinaria"],
        2,
        "Projeto conceitual de organização de conhecimento veterinário com IA e dados.",
        """
# Córtex Digital

Córtex Digital é uma frente conceitual ligada aos projetos veterinários de IA e dados. O nome representa a ideia de criar uma camada inteligente de organização do conhecimento: uma base capaz de conectar doença, sinais, fisiopatologia, conduta, classificação, vocabulário controlado e raciocínio clínico.

Ele funciona como síntese entre os objetivos do DiagnósIA, VetDex e VDC: estruturar informação veterinária para que sistemas futuros consigam consultar, explicar, comparar, filtrar e apoiar decisões com maior rastreabilidade.
""",
    ),
]

for path, title, category, tags, priority, summary, body in docs:
    write_md(path, title, category, tags, priority, summary, body)

write_md(
    "knowledge/projetos/ebookgenerator/overview.md",
    "Ebook Generator",
    "projetos",
    ["ebook-generator", "n8n", "llm", "openai", "gamma", "automacao-editorial"],
    1,
    "Projeto autoral de automação editorial com LLMs, n8n, OpenAI, Gamma e JSON estruturado.",
    f"""
# Ebook Generator

O Ebook Generator é um workflow autoral em n8n para transformar briefing e dados agregados em produto editorial premium. Ele usa prompts encadeados, OpenAI API, Gamma API, JSON estruturado, parsing, normalização, controle de custo, construção de product brain e memória conversacional.

## O que o workflow faz

- Recebe tema, briefing ou dataset consolidado.
- Planeja estratégia editorial, posicionamento, público, promessa central e direção visual.
- Pesquisa padrões de comunidades, relatos e sinais de mercado.
- Sintetiza dores, linguagem, objeções, oportunidades e ângulo editorial.
- Gera blueprint estratégico em JSON.
- Aciona um escritor editorial especializado para produzir a obra final estruturada.
- Monta payload para Gamma com formato, dimensões, estilo visual, densidade de texto, quantidade de imagens e instruções editoriais.
- Cria/atualiza product brain reutilizável para agentes conversacionais e memória de produto.
- Integra Supabase para persistência de ebook, usuário, eventos de conversa e memória.

## Nós principais identificados no JSON

{bulletize(ebook_nodes[:28])}

## Competências demonstradas

Esse projeto demonstra engenharia de prompts, automação com n8n, uso de APIs, JSON estruturado, integração com OpenAI/Gamma/Supabase, controle de tokens/custo, tratamento de erro, pipeline editorial e pensamento de produto. O valor não está só em gerar texto, mas em transformar um processo editorial inteiro em arquitetura repetível.
""",
)

write_md(
    "knowledge/projetos/autobot/overview.md",
    "Autobot - Agente pessoal com RAG e voz",
    "projetos",
    ["autobot", "rag", "voz", "fastapi", "react", "langchain", "chroma"],
    1,
    "Projeto do portfólio conversacional com RAG, voz, relatórios estáticos e pacote recrutador.",
    """
# Autobot - Agente pessoal com RAG e voz

Autobot é o projeto deste portfólio conversacional: um agente pessoal que simula uma entrevista comigo para recrutadores, respondendo em primeira pessoa com base em uma base Markdown pública e curada.

## Stack

- Backend em Python, FastAPI e LangChain.
- Banco vetorial ChromaDB local.
- OpenAI API para chat, embeddings, Whisper e TTS.
- Frontend em React/Vite.
- Interface com mapa mental, chat, voz opcional, relatórios por nó, galeria pessoal e botão Extrair Gabriel.

## Decisões de produto

O projeto separa conteúdo editorial público da lógica do agente. Os Markdown ficam em `/knowledge`, a ingestão é manual, o Chroma é recriado de forma limpa e o frontend mostra respostas sem poluir a UI com fontes. Há relatórios estáticos por nó para consulta rápida sem gastar token.

## O que o projeto demonstra

Autobot mostra capacidade de arquitetar uma aplicação full stack com IA, RAG, UX, tratamento de erro, voz, segurança de chave, API REST, pacotes de download e deploy planejado. Também mostra uma preocupação importante: quando o agente não tem contexto suficiente, ele deve admitir limite em vez de inventar.
""",
)

reports = {
    "gabriel": "# Gabriel Camim Santos\n\nSou Gabriel Camim Santos, profissional de tecnologia focado em IA generativa, dados, automação, RevOps e arquitetura de sistemas. Minha base vem do técnico em Informática, da experiência com sistemas críticos e de projetos autorais em DGE, Ebook Generator, Autobot e IA aplicada à Medicina Veterinária.\n\nO que eu levo para uma vaga é a combinação entre operação real e construção. Já atuei em suporte a ERP, PDV, infraestrutura hospitalar e ambientes de produção; ao mesmo tempo, crio projetos com LLMs, n8n, APIs, SQL, PostgreSQL, React, FastAPI, LangChain, Chroma, OpenAI, Vercel e GitHub.\n\nTenho perfil autodidata, analítico e construtor. Gosto de pegar problemas ambíguos e transformá-los em fluxo, modelo, regra, banco, API, interface ou documentação executável.\n\nCuriosidades relevantes: aprendi a ler e escrever com 4 anos, passei na ETEC com a segunda maior nota da instituição, conquistei bolsa 100% no ProUni para Medicina Veterinária sem abrir um livro de estudo e tenho 14 tatuagens, sendo 12 de animais.\n\nMeus objetivos são me formar em Data Science e criar alguma tecnologia nova, ou um sistema/app globalmente utilizado.",
    "trajetoria": "# Trajetória\n\nMinha trajetória começa na base técnica em Informática pela ETEC Adhemar Batista Heméritas, entre 2018 e 2020. Entre 2021 e 2023, trabalhei pela Spread Tecnologia no Hospital Israelita Albert Einstein, em ambiente de alta criticidade. Entre 2023 e 2025, cursei Medicina Veterinária na UNISA e conectei esse domínio a IA aplicada, dados clínicos e classificação de doenças. Em 2025, atuei na Software Show / Cacau Show com suporte aos sistemas Master Retail PDV e Master Retail Desktop. A partir de 2026, consolidei projetos autorais em dados, IA e automação: DGE, Ebook Generator e Autobot.",
    "projetos": "# Projetos\n\n## DGE - Data Growth Engine\n\nSistema autoral para governança e implantação de e-commerce orientado por dados. Estrutura uma camada de inteligência operacional auditável para conectar ERP, Bling, marketplaces, produtos, SKUs, pedidos, estoque, frete, BI, KPIs, fórmulas, snapshots, lineage, variâncias e reforecast governado.\n\n## Ebook Generator\n\nWorkflow em n8n com OpenAI API, Gamma API, JSON estruturado e Supabase para transformar briefing em blueprint editorial, obra final, payload visual e product brain reutilizável.\n\n## Autobot\n\nPortfólio conversacional com RAG e voz, usando FastAPI, LangChain, ChromaDB, OpenAI, Whisper, TTS e React.\n\n## DiagnósIA, VetDex, VDC e Córtex Digital\n\nProjetos de IA e dados aplicados à Medicina Veterinária, com foco em classificação multieixo, banco relacional, vocabulários controlados, dados clínicos e apoio à decisão.",
    "stack": "# Stack\n\nMinha stack combina IA generativa, dados, automação, backend e produto.\n\nEm IA e LLMs, trabalho com OpenAI, Claude, LangChain, engenharia de prompts, agentes, structured outputs, tool calling, RAG, memória, guardrails e avaliação de respostas.\n\nEm dados, uso SQL, PostgreSQL, modelagem relacional, schemas, KPIs, forecast/reforecast, governança, lineage, documentação de métricas, Power BI e Superset.\n\nEm automação, uso n8n, APIs REST, webhooks, JSON, OpenAI API, Gamma API, parsing, tratamento de erros e pipelines de dados textuais.\n\nEm desenvolvimento, uso JavaScript, Python, HTML/CSS, GitHub, VS Code, Vercel, Docker, FastAPI, React, backend, APIs internas, autenticação, permissões, CI/CD e banco multitenant.",
    "experiencia": "# Experiência\n\nMinha experiência formal passa por suporte a sistemas críticos, ERP, PDV, infraestrutura e operação.\n\nNa Software Show / Cacau Show, atuei com Master Retail PDV e Desktop em franquias, lidando com estabilidade de loja, emissão fiscal, pagamentos, retaguarda, frente de caixa, TEF, SAT, pinpads, impressoras térmicas, totens e chamados de produção.\n\nNa Spread Tecnologia / Hospital Israelita Albert Einstein, atuei em ambiente hospitalar crítico, com rollout, Field Service, Service Desk N2 remoto, Active Directory, SCCM, Cisco ACS, Cisco ISE, ITSM, instalação de sistemas, configuração de estações Windows e automação em `.bat` adotada pela equipe.\n\nO aprendizado principal dessas experiências é simples: tecnologia precisa funcionar com usuário real, urgência real e impacto real.",
    "mercado": "# Mercado\n\nMinha visão de mercado é que IA, dados e automação só geram valor quando viram estrutura operacional. Não basta chamar API de LLM; é preciso saber onde o dado entra, como ele é validado, qual decisão apoia, qual risco reduz e como o sistema preserva contexto.\n\nEu gosto de problemas em que há muita informação dispersa, processo manual, decisão recorrente, baixa rastreabilidade ou dificuldade de transformar conhecimento em operação. É nesse tipo de cenário que RAG, workflows, bancos relacionais, agentes, dashboards e automações podem ter impacto real.",
    "entrevista": "# Entrevista\n\nEm entrevista, eu me apresento como alguém que conecta operação real com construção técnica. Tenho experiência em suporte e sistemas críticos, mas meu diferencial atual está em projetos autorais de IA, dados, automação e arquitetura.\n\nSe perguntarem por que eu deveria ser contratado, minha resposta é: porque eu aprendo rápido, transformo problema nebuloso em estrutura e consigo conversar tanto com a camada técnica quanto com o negócio.\n\nSe perguntarem sobre limitações, eu respondo com honestidade: ainda estou consolidando minha formação formal em Data Science e amadurecendo em produção de software em escala, mas compenso isso com estudo intenso, documentação, projetos próprios e disposição para aprender com times mais experientes.\n\nExperiências marcantes: automação `.bat` criada no Einstein, DGE como arquitetura de inteligência operacional, Ebook Generator como pipeline de automação editorial e VetDex/VDC como modelagem de conhecimento veterinário.",
    "materiais": "# Materiais\n\nO botão Extrair Gabriel baixa o pacote recrutador com documentos curados para avaliação rápida: currículo, perfil técnico, carta de motivação e resumo de projetos.\n\nO conteúdo é público e preparado para recrutadores. Chaves de API, `.env`, Chroma local, caches e arquivos temporários não entram no pacote.",
}

for key, body in reports.items():
    write_md(
        f"knowledge/reports/{key}.md",
        f"Relatório {key.title()}",
        "reports",
        ["relatorio", key],
        1,
        f"Relatório estático do nó {key}.",
        body,
    )

write_md(
    "knowledge/entrevista/faq.md",
    "FAQ de entrevista",
    "entrevista",
    ["faq", "entrevista", "recrutamento"],
    1,
    "Perguntas frequentes de entrevista e respostas baseadas no perfil do Gabriel.",
    """
# FAQ de entrevista

## Quem é Gabriel Camim Santos?

Sou um profissional de tecnologia focado em IA generativa, dados, automação, RevOps e arquitetura de sistemas. Tenho experiência em suporte a sistemas críticos e desenvolvo projetos autorais que conectam LLMs, dados, APIs, automações e produto.

## Qual é seu diferencial?

Meu diferencial é conseguir unir vivência operacional com construção técnica. Eu entendo que sistema precisa funcionar em produção, mas também gosto de modelar dados, desenhar arquitetura, documentar raciocínio e transformar ambiguidade em solução.

## Quais projetos melhor representam seu perfil?

DGE, Ebook Generator, Autobot e os projetos veterinários VetDex/VDC/DiagnósIA. Eles mostram minha capacidade de estruturar sistemas com dados, IA, automação e domínio de negócio.

## O que você ainda está desenvolvendo?

Estou consolidando minha formação formal em Data Science e amadurecendo meu repertório de engenharia de software em escala. Ao mesmo tempo, venho compensando isso com projetos práticos, documentação e estudo contínuo.
""",
)

write_md(
    "knowledge/entrevista/experiencias-marcantes.md",
    "Experiências marcantes para entrevista",
    "entrevista",
    ["experiencias-marcantes", "entrevista", "casos"],
    1,
    "Casos e experiências marcantes para contar em entrevistas.",
    """
# Experiências marcantes

## Automação adotada no Einstein

Durante minha atuação no Hospital Israelita Albert Einstein, criei uma automação em `.bat` para executar instalações e configurações padronizadas. A automação foi adotada pela equipe de suporte, reduzindo esforço repetitivo e acelerando o fluxo operacional.

## DGE como arquitetura de inteligência operacional

A DGE é um exemplo forte de como eu penso sistemas: dados entram com governança, viram KPIs, alimentam fórmulas, preservam histórico, mostram variâncias e sustentam decisão humana.

## Ebook Generator como automação ponta a ponta

O Ebook Generator mostra minha capacidade de orquestrar LLMs, n8n, APIs, JSON, parsing, tratamento de erro, Gamma, Supabase, product brain e memória em um workflow de produto real.

## VetDex/VDC como modelagem de conhecimento

VetDex e VDC representam minha habilidade de mergulhar em um domínio complexo, como Medicina Veterinária, e traduzir conhecimento em classificação, banco relacional, vocabulários, eixos e sistemas consultáveis.
""",
)

MATERIALS.mkdir(parents=True, exist_ok=True)
shutil.copy2(CV_DOCX, MATERIALS / "Curriculo_Gabriel_Camim_Dados_IA_Automacao.docx")

(MATERIALS / "curriculo.md").write_text(
    """# Currículo - Gabriel Camim Santos

## Contato

São Paulo - SP | camim2003@gmail.com | +55 (11) 95804-8353 | LinkedIn: Gabriel Camim

## Resumo

Profissional de tecnologia com formação técnica em Informática, experiência em suporte a sistemas críticos, ERP, PDV e ambientes de produção. Atua na interseção entre IA generativa, dados, automação, RevOps e arquitetura de sistemas.

## Projetos técnicos autorais

- DGE - Data Growth Engine: inteligência operacional auditável para e-commerce, KPIs, projeções, Bling, ERP, estoque, pedidos, frete e reforecast governado.
- Ebook Generator: workflow em n8n com OpenAI, Gamma, JSON estruturado, product brain, memória e automação editorial.
- DiagnósIA, VetDex, VDC e Córtex Digital: IA, dados clínicos, classificação multieixo e sistemas de apoio à decisão em Medicina Veterinária.
- Autobot: portfólio conversacional com RAG, voz, FastAPI, LangChain, ChromaDB e React.

## Experiência

- Software Show / Cacau Show - Analista de Suporte Jr, Jan/2025.
- Spread Tecnologia / Hospital Israelita Albert Einstein - Analista de Suporte Jr, Ago/2021 a Fev/2023.

## Formação

- Data Science - FIAP, previsão Jul/2026 a Dez/2028.
- Medicina Veterinária - UNISA, Fev/2023 a Set/2025, incompleta.
- Ensino Médio Técnico em Informática - ETEC Adhemar Batista Heméritas, 2018 a 2020.

## Idiomas

Inglês com leitura fluente, escrita avançada e conversação avançada.
""",
    encoding="utf-8",
    newline="\n",
)

(MATERIALS / "perfil-tecnico.md").write_text(
    """# Perfil técnico - Gabriel Camim Santos

Gabriel atua na interseção entre IA generativa, dados, automação, RevOps e arquitetura de sistemas. Tem experiência com suporte a sistemas críticos e projetos autorais usando LLMs, n8n, APIs, JSON estruturado, SQL, PostgreSQL, FastAPI, LangChain, ChromaDB, React, Vercel, GitHub e OpenAI.

Competências principais:

- IA generativa, agentes, RAG, tool calling, prompts e structured outputs.
- Dados, SQL, PostgreSQL, modelagem relacional, KPIs, lineage e forecast/reforecast.
- Automação com n8n, APIs REST, webhooks, parsing e tratamento de erro.
- Desenvolvimento full stack com Python, JavaScript, FastAPI, React e arquitetura modular.
- RevOps, CRM, funil comercial, B2B/B2C, lead scoring e inteligência de crescimento.
""",
    encoding="utf-8",
    newline="\n",
)

(MATERIALS / "resumo-projetos.md").write_text(
    """# Resumo de projetos - Gabriel Camim Santos

## DGE - Data Growth Engine

Sistema autoral de inteligência operacional auditável para e-commerce, conectando ERP, Bling, marketplaces, estoque, pedidos, SKUs, frete, BI, KPIs, fórmulas, snapshots, lineage, variâncias e reforecast governado.

## Ebook Generator

Workflow em n8n com OpenAI, Gamma, JSON estruturado, Supabase e memória para transformar briefing em blueprint editorial, obra final, payload visual e product brain reutilizável.

## Autobot

Agente pessoal com RAG e voz para portfólio interativo, usando FastAPI, LangChain, ChromaDB, OpenAI, Whisper, TTS e React.

## DiagnósIA, VetDex, VDC e Córtex Digital

Projetos de IA e dados aplicados à Medicina Veterinária, com foco em classificação multieixo, banco relacional, vocabulários controlados, filtros e apoio à decisão.
""",
    encoding="utf-8",
    newline="\n",
)

(MATERIALS / "README.md").write_text(
    """# Pacote recrutador - Extrair Gabriel

Este pacote reúne documentos públicos e curados para avaliação rápida de Gabriel Camim Santos.

Arquivos principais:

- `Curriculo_Gabriel_Camim_Dados_IA_Automacao.docx`
- `curriculo.md`
- `perfil-tecnico.md`
- `resumo-projetos.md`
- `carta-motivacao.md`

Nenhum segredo, chave de API, `.env`, cache ou índice local entra neste pacote.
""",
    encoding="utf-8",
    newline="\n",
)

import_dir = KNOWLEDGE / "projetos" / "dge" / "imported"
resolved_import = import_dir.resolve()
if resolved_import.exists():
    if ROOT.resolve() not in resolved_import.parents:
        raise RuntimeError(f"Unsafe import directory: {resolved_import}")
    shutil.rmtree(resolved_import)
import_dir.mkdir(parents=True, exist_ok=True)

skip_parts = {".git", "node_modules", ".next", "dist", "build", "__pycache__"}
secret_patterns = [
    (re.compile(r"sk-[A-Za-z0-9_\-]{20,}"), "[SECRET_REDACTED]"),
    (re.compile(r"(?i)(OPENAI_API_KEY\s*=\s*)\S+"), r"\1[SECRET_REDACTED]"),
    (re.compile(r"(?i)(GAMMA_API_KEY\s*=\s*)\S+"), r"\1[SECRET_REDACTED]"),
]


def sanitize(text: str) -> str:
    output = text
    for pattern, replacement in secret_patterns:
        output = pattern.sub(replacement, output)
    return output


mds = sorted(
    path
    for path in DGE_ROOT.rglob("*.md")
    if path.is_file() and not any(part in skip_parts for part in path.parts)
)

for source in mds:
    rel = source.relative_to(DGE_ROOT).as_posix()
    raw = sanitize(source.read_text(encoding="utf-8", errors="replace"))
    title_match = re.search(r"^#\s+(.+)$", raw, flags=re.MULTILINE)
    title = title_match.group(1).strip() if title_match else source.stem.replace("-", " ").title()
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", rel).strip("-").lower()
    tags = ["dge", "fonte-original"]
    if "architecture" in rel:
        tags.append("arquitetura")
    if "contract" in rel:
        tags.append("contratos")
    if "implementation" in rel:
        tags.append("implementacao")
    body = f"# {title}\n\nFonte original DGE 2.0: `{rel}`.\n\n---\n\n{raw.strip()}"
    write_md(
        import_dir / f"{slug}.md",
        f"DGE fonte - {title}",
        "projetos",
        tags,
        3,
        f"Documento original importado da base DGE 2.0: {rel}.",
        body,
        source_path=rel,
    )

print(f"Markdown atualizados. DGE importados: {len(mds)}. Currículo copiado para pacote recrutador.")
